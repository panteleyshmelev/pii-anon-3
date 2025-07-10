# app/routes/documents.py
import uuid
import json
import traceback
from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import FileResponse, PlainTextResponse
from app.core.masking import get_pii_entities, create_masked_text_and_map, PiiEntity # Import new functions
import os
import fitz  # PyMuPDF for text extraction from PDF
import docx  # For .docx files
# collections.defaultdict is no longer needed here, it's in masking.py

router = APIRouter()

UPLOAD_DIR = "data/uploads"
MASKED_TXTS_DIR = "data/masked_txts"
UNMASKED_TXTS_DIR = "data/unmasked_txts"
MAPPING_FILE = "data/masking_maps.json" # This will store {file_id: unmasking_map_for_file}

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(MASKED_TXTS_DIR, exist_ok=True)
os.makedirs(UNMASKED_TXTS_DIR, exist_ok=True)

def load_mappings():
    if os.path.exists(MAPPING_FILE):
        try:
            with open(MAPPING_FILE, "r", encoding="utf-8") as f:
                content = f.read()
                if not content: return {}
                return json.loads(content)
        except json.JSONDecodeError: return {}
    return {}

def save_mappings(mappings):
    with open(MAPPING_FILE, "w", encoding="utf-8") as f:
        json.dump(mappings, f, indent=4)

@router.post("/process-to-masked-txt", tags=["Text Processing"])
async def process_to_masked_txt(file: UploadFile = File(...)):
    file_id = str(uuid.uuid4())
    original_file_path = os.path.join(UPLOAD_DIR, f"{file_id}_{file.filename}")
    masked_txt_path = os.path.join(MASKED_TXTS_DIR, f"masked_{file_id}.txt")

    with open(original_file_path, "wb") as buffer:
        buffer.write(await file.read())

    full_text_parts = []
    file_extension = file.filename.lower().split('.')[-1] if '.' in file.filename else ''

    try:
        if file_extension == "pdf":
            doc_pdf = fitz.open(original_file_path)
            for page in doc_pdf:
                full_text_parts.append(page.get_text())
            doc_pdf.close()
        elif file_extension == "txt":
            with open(original_file_path, "r", encoding="utf-8") as f_txt:
                full_text_parts.append(f_txt.read())
        elif file_extension == "docx":
            try:
                doc_word = docx.Document(original_file_path)
                for para in doc_word.paragraphs:
                    full_text_parts.append(para.text)
                for table in doc_word.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            full_text_parts.append(cell.text)
            except Exception as e_docx:
                raise HTTPException(status_code=400, detail=f"Error processing .docx file. Ensure it's a valid .docx format. (Details: {str(e_docx)})")
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: .{file_extension}. Please upload PDF, TXT, or DOCX.")

        full_text = "\n".join(full_text_parts)

        if not full_text.strip():
            raise HTTPException(status_code=400, detail="No text content found or extracted from the uploaded file.")

    except HTTPException as http_exc:
        if os.path.exists(original_file_path): os.remove(original_file_path)
        raise http_exc
    except Exception as e:
        if os.path.exists(original_file_path): os.remove(original_file_path)
        raise HTTPException(status_code=500, detail=f"Error extracting text from file '{file.filename}': {str(e)}")

    try:
        pii_entities: List[PiiEntity] = get_pii_entities(full_text)
        
        if not pii_entities:
            masked_text = full_text
            unmasking_map_for_file = {}
            detail_message = "File processed. No PII found by LLM. Original text saved."
        else:
            masked_text, unmasking_map_for_file = create_masked_text_and_map(full_text, pii_entities)
            detail_message = f"File '{file.filename}' processed, text extracted, PII masked, and saved as .txt."

    except Exception as e:
        print("\n--- DEBUG: AN UNEXPECTED ERROR OCCURRED DURING PII PROCESSING. FULL TRACEBACK: ---")
        traceback.print_exc()
        print("-----------------------------------------------------------\n")
        if os.path.exists(original_file_path): os.remove(original_file_path)
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred during PII processing. Check server logs for traceback.")

    # --- Save Unmasking Map ---
    all_document_mappings = load_mappings()
    all_document_mappings[file_id] = unmasking_map_for_file
    save_mappings(all_document_mappings)

    # --- START: DEBUG PRINTS FOR FILE SAVING ---
    print("\n--- DEBUG: Preparing to save masked text file. ---")
    print(f"--- DEBUG: File path it will be saved to: {masked_txt_path}")
    print(f"--- DEBUG: Length of masked text to be written: {len(masked_text)} chars.")
    # You can uncomment the next line to see the first 500 characters of the masked text
    # print(f"--- DEBUG: Masked text preview: {masked_text[:500]}")
    print("-------------------------------------------------\n")
    # --- END: DEBUG PRINTS FOR FILE SAVING ---

    # --- Save Masked Text to .txt file ---
    try:
        with open(masked_txt_path, "w", encoding="utf-8") as f_out:
            f_out.write(masked_text)
        print("--- DEBUG: File write operation completed successfully. ---")
    except Exception as e:
        print(f"--- DEBUG: ERROR occurred during file write operation: {e} ---")
        # Still raise an error to the client if writing fails
        raise HTTPException(status_code=500, detail=f"Failed to save masked file. Error: {str(e)}")


    if os.path.exists(original_file_path):
        os.remove(original_file_path)

    return {
        "file_id": file_id,
        "original_filename": file.filename,
        "masked_txt_file_path": masked_txt_path,
        "detail": detail_message
    }

@router.get("/demask-txt/{file_id}", tags=["Text Processing"])
async def demask_txt(file_id: str):
    masked_txt_path = os.path.join(MASKED_TXTS_DIR, f"masked_{file_id}.txt")
    unmasked_txt_path = os.path.join(UNMASKED_TXTS_DIR, f"unmasked_{file_id}.txt") # For saving demasked text

    if not os.path.exists(masked_txt_path):
        raise HTTPException(status_code=404, detail=f"Masked text file not found for ID: {file_id}")

    all_document_mappings = load_mappings()
    if file_id not in all_document_mappings:
        raise HTTPException(status_code=404, detail=f"Unmasking map not found for ID: {file_id}")

    unmasking_map_for_file = all_document_mappings[file_id] # This is {placeholder: original_text}

    with open(masked_txt_path, "r", encoding="utf-8") as f_in:
        masked_text_content = f_in.read()

    # --- Demasking the Text (String Replacement) ---
    unmasked_text = masked_text_content
    # Replace longest placeholders first to avoid issues if placeholders could be substrings of others (unlikely with current format)
    sorted_placeholders = sorted(unmasking_map_for_file.keys(), key=len, reverse=True)
    
    for placeholder in sorted_placeholders:
        original_value = unmasking_map_for_file[placeholder]
        unmasked_text = unmasked_text.replace(placeholder, original_value)

    # --- Save Unmasked Text to .txt file (optional, or just return it) ---
    with open(unmasked_txt_path, "w", encoding="utf-8") as f_out:
        f_out.write(unmasked_text)
        
    return PlainTextResponse(content=unmasked_text)